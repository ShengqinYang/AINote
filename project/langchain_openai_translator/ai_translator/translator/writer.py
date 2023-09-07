import os
from reportlab.lib import colors, pagesizes, units
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
)
from docx import Document
from docx.shared import Pt, Inches

from book import Book, ContentType
from utils import LOG


class Writer:
    def __init__(self):
        pass

    def save_translated_book(self, book: Book, ouput_file_format: str, translation_style: str):
        LOG.debug(ouput_file_format)

        if ouput_file_format.lower() == "pdf":
            file_suffix = f'_translated_{translation_style}.pdf'
            output_file_path = self._save_translated_book_pdf(book, file_suffix)
        elif ouput_file_format.lower() == "markdown":
            file_suffix = f'_translated_{translation_style}.md'
            output_file_path = self._save_translated_book_markdown(book, file_suffix)
        elif ouput_file_format.lower() == "word":
            file_suffix = f'_translated_{translation_style}.docx'
            output_file_path = self._save_translated_book_word(book, file_suffix)
        else:
            LOG.error(f"不支持文件类型: {ouput_file_format}")
            return ""

        LOG.info(f"翻译完成，文件保存至: {output_file_path}")

        return output_file_path

    def _save_translated_book_pdf(self, book: Book, file_suffix: str, output_file_path: str = None):
        output_file_path = book.pdf_file_path.replace('.pdf', file_suffix)
        LOG.info(f"开始导出: {output_file_path}")

        # Register Chinese font
        font_path = "../fonts/simsun.ttc"  # 请将此路径替换为您的字体文件路径
        pdfmetrics.registerFont(TTFont("SimSun", font_path))

        # Create a new ParagraphStyle with the SimSun font
        simsun_style = ParagraphStyle('SimSun', fontName='SimSun', fontSize=12, leading=14)

        # Create a PDF document
        doc = SimpleDocTemplate(output_file_path, pagesize=pagesizes.letter)
        styles = getSampleStyleSheet()
        story = []

        # Iterate over the pages and contents
        for page in book.pages:
            for content in page.contents:
                if content.status:
                    if content.content_type == ContentType.TEXT:
                        # Add translated text to the PDF
                        text = content.translation
                        para = Paragraph(text, simsun_style)
                        story.append(para)

                    elif content.content_type == ContentType.TABLE:
                        # Add table to the PDF
                        table = content.translation
                        table_style = TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'SimSun'),  # 更改表头字体为 "SimSun"
                            ('FONTSIZE', (0, 0), (-1, 0), 14),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('FONTNAME', (0, 1), (-1, -1), 'SimSun'),  # 更改表格中的字体为 "SimSun"
                            ('GRID', (0, 0), (-1, -1), 1, colors.black)
                        ])
                        pdf_table = Table(table.values.tolist())
                        pdf_table.setStyle(table_style)
                        story.append(pdf_table)
            # Add a page break after each page except the last one
            if page != book.pages[-1]:
                story.append(PageBreak())

        # Save the translated book as a new PDF file
        doc.build(story)
        return output_file_path

    def _save_translated_book_markdown(self, book: Book, file_suffix: str, output_file_path: str = None):
        output_file_path = book.pdf_file_path.replace('.pdf', file_suffix)

        LOG.info(f"开始导出: {output_file_path}")
        with open(output_file_path, 'w', encoding='utf-8') as output_file:
            # Iterate over the pages and contents
            for page in book.pages:
                for content in page.contents:
                    if content.status:
                        if content.content_type == ContentType.TEXT:
                            # Add translated text to the Markdown file
                            text = content.translation
                            output_file.write(text + '\n\n')

                        elif content.content_type == ContentType.TABLE:
                            # Add table to the Markdown file
                            table = content.translation
                            header = '| ' + ' | '.join(str(column) for column in table.columns) + ' |' + '\n'
                            separator = '| ' + ' | '.join(['---'] * len(table.columns)) + ' |' + '\n'
                            # body = '\n'.join(['| ' + ' | '.join(row) + ' |' for row in table.values.tolist()]) + '\n\n'
                            body = '\n'.join(['| ' + ' | '.join(str(cell) for cell in row) + ' |' for row in
                                              table.values.tolist()]) + '\n\n'
                            output_file.write(header + separator + body)

                # Add a page break (horizontal rule) after each page except the last one
                if page != book.pages[-1]:
                    output_file.write('---\n\n')

        return output_file_path

    def _save_translated_book_word(self, book: Book, file_suffix: str, output_file_path: str = None):
        output_file_path = book.pdf_file_path.replace('.pdf', file_suffix)
        LOG.info(f"开始导出: {output_file_path}")


        doc = Document()

        # Iterate over the pages and contents
        for page in book.pages:
            for content in page.contents:
                if content.status:
                    if content.content_type == ContentType.TEXT:
                        # Add translated text to the Word document
                        text = content.translation
                        p = doc.add_paragraph()
                        run = p.add_run(text)
                        run.font.size = Pt(12)  # Set font size to 12pt
                        run.font.name = 'Arial'  # Set font to Arial

                    elif content.content_type == ContentType.TABLE:
                        # Add table to the Word document
                        table = content.translation
                        rows = len(table.values)
                        cols = len(table.columns)
                        doc.add_table(rows=rows + 1, cols=cols)
                        table_cells = doc.tables[-1].rows
                        header_row = table_cells[0].cells
                        for i, column in enumerate(table.columns):
                            header_row[i].text = str(column)
                        for i, row in enumerate(table.values.tolist()):
                            table_row = table_cells[i + 1].cells
                            for j, cell in enumerate(row):
                                table_row[j].text = str(cell)
                    elif content.content_type == ContentType.IMAGE:
                        # Add image to the Word document
                        image_path = content.translation
                        doc.add_picture(image_path, width=Inches(3))  # Adjust width as needed

            # Add a page break after each page except the last one
            if page != book.pages[-1]:
                doc.add_page_break()

        doc.save(output_file_path)

        LOG.info(f"翻译完成: {output_file_path}")
        return output_file_path
